#!/usr/bin/env python3
"""Convert the simulation data request document to OnlyOffice-compatible docx format - v3 professional."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading if any
    for child in list(tcPr):
        if child.tag.endswith('shd'):
            tcPr.remove(child)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    shading_elm.set(qn('w:val'), 'clear')
    tcPr.append(shading_elm)


def set_cell_vertical_alignment(cell, alignment='center'):
    """Set cell vertical alignment."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment)
    tcPr.append(vAlign)


def create_styled_table(doc, data, col_widths_cm=None, font_size=9):
    """Create a professionally styled table."""
    num_rows = len(data)
    num_cols = len(data[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '6')  # Border width
        element.set(qn('w:color'), '666666')  # Gray border
        tblBorders.append(element)
    
    # Remove old borders if exist
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    tblPr.append(tblBorders)
    
    # Set column widths
    if col_widths_cm:
        for i, width_cm in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(width_cm)
    
    # Populate table
    for i, row_data in enumerate(data):
        row = table.rows[i]
        # Set minimum row height
        row.height = Cm(0.8)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            
            # Set vertical alignment
            set_cell_vertical_alignment(cell, 'center')
            
            # Add text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_data))
            run.font.name = 'Microsoft YaHei'  # 微软雅黑 - more modern
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            # Header row styling
            if i == 0:
                set_cell_shading(cell, '2F5496')  # Dark blue
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(font_size + 1)
            else:
                # Alternate row colors
                if i % 2 == 0:
                    set_cell_shading(cell, 'E7E6E6')  # Light gray
                else:
                    set_cell_shading(cell, 'FFFFFF')  # White
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    return table


def create_document():
    doc = Document()
    
    # ===== Setup Styles =====
    styles = doc.styles
    
    # Modify Normal style - use 微软雅黑 for better rendering
    style_normal = styles['Normal']
    style_normal.font.name = 'Microsoft YaHei'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Remove underlines from all heading styles
    for heading_name in ['Heading 1', 'Heading 2', 'Heading 3']:
        style = styles[heading_name]
        style.font.name = 'Microsoft YaHei'
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.underline = False
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Heading 1 - 一级标题
    styles['Heading 1'].font.size = Pt(16)
    styles['Heading 1'].font.bold = True
    styles['Heading 1'].paragraph_format.space_before = Pt(24)
    styles['Heading 1'].paragraph_format.space_after = Pt(12)
    
    # Heading 2 - 二级标题
    styles['Heading 2'].font.size = Pt(14)
    styles['Heading 2'].font.bold = True
    styles['Heading 2'].paragraph_format.space_before = Pt(18)
    styles['Heading 2'].paragraph_format.space_after = Pt(6)
    
    # Heading 3 - 三级标题
    styles['Heading 3'].font.size = Pt(12)
    styles['Heading 3'].font.bold = True
    styles['Heading 3'].paragraph_format.space_before = Pt(12)
    styles['Heading 3'].paragraph_format.space_after = Pt(6)
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21)  # A4
        section.page_height = Cm(29.7)
    
    # ===== TITLE =====
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run('KYN28高压开关柜多物理场仿真分析')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_subtitle.add_run('基础数据与条件输入书')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p_subtitle.paragraph_format.space_after = Pt(24)
    
    # Header info table
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ('编制单位', '仿真分析工作组'),
        ('适用对象', '甲方技术中心 / 设计研发部'),
        ('日    期', '2025年12月12日'),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        
        cell_label.text = ''
        p = cell_label.paragraphs[0]
        run = p.add_run(f'{label}：')
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(11)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        cell_value.text = ''
        p = cell_value.paragraphs[0]
        run = p.add_run(value)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
    
    doc.add_paragraph()
    
    # ===== SECTION 1 =====
    doc.add_paragraph('1. 总则', style='Heading 1')
    
    doc.add_paragraph('1.1 目的', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('为确保KYN28高压开关柜在电磁场、温度场及机械动力学等多物理场仿真中的计算精度，真实反映产品在实际工况下的物理性能，特编制本数据需求书。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('1.2 适用范围', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('本文档涵盖了仿真所需的核心几何模型、材料物理属性、热工边界条件及动力学参数。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('1.3 重要性说明', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('仿真结果的置信度直接取决于输入数据的准确性。若关键参数缺失或使用通用估算值，可能导致电场击穿风险评估失效或温升计算偏差超过20%。请甲方依据本文件逐项确认。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    # ===== SECTION 2 =====
    doc.add_paragraph('2. 几何数模需求 (Geometric Data)', style='Heading 1')
    
    doc.add_paragraph('2.1 原始三维装配设计文件', style='Heading 2')
    doc.add_paragraph('2.1.1 需求内容', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('请提供SolidWorks (.sldasm)、Inventor (.iam) 或 Creo (.asm) 原生格式的三维装配图。仅在无法提供原生文件时，接受STEP (.stp) 格式。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('2.1.2 技术必要性', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('原生参数化文件支持特征抑制（Suppress）与几何简化（如去除螺栓、倒角），是高效前处理的基础。中间格式（SAT/IGES）往往丢失装配树，导致模型清理周期大幅延长。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('2.1.3 缺失风险', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('若采用现有SAT导出模型，因从属于死实体且包含大量非分析特征，预计网格划分失败率极高，且无法进行参数化优化设计。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('2.2 真空灭弧室 (VI) 内部结构', style='Heading 2')
    doc.add_paragraph('2.2.1 需求内容', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('需提供灭弧室内部结构的2D剖面图（DWG/PDF）或简化的3D几何体。重点包含：')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    for i, item in enumerate(['动、静触头的几何型面（如杯状磁场触头或纵磁触头）；',
                              '主屏蔽罩与辅助屏蔽罩的坐标位置；',
                              '波纹管的大致包络尺寸。'], 1):
        p = doc.add_paragraph()
        p.add_run(f'({i}) {item}')
        p.paragraph_format.left_indent = Cm(1.5)
    
    doc.add_paragraph('2.2.2 技术必要性', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('触头型面的曲率半径直接决定真空间隙的最大电场强度。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('2.2.3 缺失风险', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('若缺失此项，无法评估灭弧室内部的绝缘击穿风险及电位分布，电场仿真将失去核心价值。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('2.3 操作机构与脱扣器组件', style='Heading 2')
    doc.add_paragraph('2.3.1 需求内容', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('独立的电磁铁组件三维模型，包含线圈绑组包络、静铁芯、动铁芯（衔铁）及工作气隙。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('2.3.2 技术必要性', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('整柜模型中通常简化了机构细节。瞬态动力学仿真必须依赖精确的磁路几何来计算电磁力。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('2.4 几何参数汇总表', style='Heading 2')
    
    # Table G - with wider columns
    table_g_data = [
        ['序号', '参数名称', '符号', '单位', '数据要求', '技术说明', '推荐值', '典型范围'],
        ['G1', '触头开距', 'd_open', 'mm', '设计值', '决定击穿电压', '11', '9~13'],
        ['G2', '超程', 'd_over', 'mm', '设计值', '影响接触压力', '3.5', '3~4'],
        ['G3', '触头半径', 'Rc', 'mm', '剖面图', '决定电场集中', 'R3', 'R2~R4'],
        ['G4', '气隙长度', 'δ', 'mm', '3D模型', '决定初始吸力', '10', '5~15'],
        ['G5', '爬电距离', 'L_creep', 'mm', '沿面路径', '安规指标', '240', '≥230'],
    ]
    create_styled_table(doc, table_g_data, [1.5, 2.0, 1.5, 1.2, 2.0, 2.5, 1.8, 1.8])
    
    doc.add_paragraph()
    
    # ===== SECTION 3 =====
    doc.add_paragraph('3. 材料物理参数 (Material Properties)', style='Heading 1')
    
    doc.add_paragraph('3.1 导电材料', style='Heading 2')
    doc.add_paragraph('3.1.1 需求内容', style='Heading 3')
    for i, item in enumerate(['母线/触头臂材质牌号（如T2紫铜）；',
                              '实测电导率（如 57 MS/m 或 98% IACS 等级）；',
                              '触头镀银层厚度及工艺标准。'], 1):
        p = doc.add_paragraph()
        p.add_run(f'({i}) {item}')
        p.paragraph_format.left_indent = Cm(1.5)
    
    doc.add_paragraph('3.1.2 敏感度分析', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('电导率每下降10%，焦耳热损耗将增加约10%，直接导致最终温升超标。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('3.2 导磁材料（柜体与机构）', style='Heading 2')
    doc.add_paragraph('3.2.1 需求内容', style='Heading 3')
    for i, item in enumerate(['柜体骨架材质：敷铝锌板、冷轧钢板或非磁性不锈钢；',
                              '磁性能参数：相对磁导率或B-H磁化曲线；',
                              '脱扣器铁芯材料牌号（如DT4电工纯铁）及其饱和磁化曲线。'], 1):
        p = doc.add_paragraph()
        p.add_run(f'({i}) {item}')
        p.paragraph_format.left_indent = Cm(1.5)
    
    doc.add_paragraph('3.2.2 敏感度分析', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('漏磁通在柜体上产生的涡流损耗是开关柜发热的重要组成部分，其大小与材料磁导率密切相关。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('3.3 绝缘材料', style='Heading 2')
    doc.add_paragraph('3.3.1 需求内容', style='Heading 3')
    for i, item in enumerate(['固封极柱（环氧树脂）：相对介电常数、介质损耗角正切；',
                              'SMC/DMC绝缘隔板：同上参数。'], 1):
        p = doc.add_paragraph()
        p.add_run(f'({i}) {item}')
        p.paragraph_format.left_indent = Cm(1.5)
    
    doc.add_paragraph('3.3.2 敏感度分析', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('复合介质下的电场分布完全取决于各材料介电常数的匹配度。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('3.4 材料参数汇总表', style='Heading 2')
    
    # Table M
    table_m_data = [
        ['序号', '参数名称', '符号', '单位', '数据要求', '技术说明', '推荐值', '典型范围'],
        ['M1', '电导率', 'σ', 'S/m', '实测值', '影响焦耳热', '5.7×10⁷', '≥5.6×10⁷'],
        ['M2', '相对磁导率', 'μr', '1', 'B-H曲线', '影响涡流损耗', '1', '1~2000'],
        ['M3', '饱和磁感', 'B_sat', 'T', '实测值', '决定电磁力上限', '2.15', '1.8~2.2'],
        ['M4', '相对介电常数', 'εr', '1', '50Hz下', '决定电位分布', '4.2', '3.5~5.5'],
        ['M5', '介质损耗', 'tanδ', '1', '50Hz下', '影响介质发热', '0.02', '0.005~0.05'],
    ]
    create_styled_table(doc, table_m_data, [1.5, 2.2, 1.3, 1.2, 2.0, 2.5, 1.8, 2.0])
    
    doc.add_paragraph()
    
    # ===== SECTION 4 =====
    doc.add_paragraph('4. 热工与流体边界条件 (Thermal & Airflow)', style='Heading 1')
    
    doc.add_paragraph('4.1 接触电阻', style='Heading 2')
    doc.add_paragraph('4.1.1 需求内容', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('主回路各搭接面、梅花触头配合面、动静触头咬合面的回路电阻实测值（微欧级）。若无实测数据，请提供相关企业标准或型式试验报告中的允许限值。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('4.1.2 重要性', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('接触电阻发热通常占总发热量的30%-50%，是热仿真中最关键的不确定性热源。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('4.2 环境与散热边界', style='Heading 2')
    doc.add_paragraph('4.2.1 需求内容', style='Heading 3')
    for i, item in enumerate(['设计海拔高度（如2000m以上需考虑散热降额）；',
                              '设计环境温度（如40℃）；',
                              '强制风冷风机型号或P-Q特定曲线（风量-静压曲线）。'], 1):
        p = doc.add_paragraph()
        p.add_run(f'({i}) {item}')
        p.paragraph_format.left_indent = Cm(1.5)
    
    doc.add_paragraph('4.3 热工参数汇总表', style='Heading 2')
    
    # Table T
    table_t_data = [
        ['序号', '参数名称', '符号', '单位', '数据要求', '技术说明', '推荐值', '典型范围'],
        ['T1', '接触电阻', 'Rc', 'μΩ', '实测值', '集中热源', '20', '10~50'],
        ['T2', '环境温度', 'T_amb', '°C', '设计值', '温升基准', '40', '40/55'],
        ['T3', '表面发射率', 'ε', '1', '表面处理', '辐射散热', '0.8', '0.1~0.9'],
        ['T4', '换热系数', 'h', 'W/(m²K)', '仿真计算', '对流散热', '10', '5~20'],
    ]
    create_styled_table(doc, table_t_data, [1.5, 2.0, 1.3, 1.8, 2.0, 2.2, 1.8, 1.8])
    
    doc.add_paragraph()
    
    # ===== SECTION 5 =====
    doc.add_paragraph('5. 动力学系统参数 (Dynamics)', style='Heading 1')
    
    doc.add_paragraph('5.1 运动质量', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('需提供动触头、导电杆、绝缘拉杆及衔铁等所有运动部件的折算总质量（kg）。用于计算动作加速度。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('5.2 机械反力特性', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('需提供触头初压力弹簧、超程弹簧及分闸弹簧的刚度系数（k）及预压缩量。电磁力必须克服此机械反力方能做功。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('5.3 驱动电路', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('脱扣线圈的匝数、线径及驱动电压波形（恒压或电容放电）。用于计算安匝数（磁动势）。')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    doc.add_paragraph('5.4 动力学参数汇总表', style='Heading 2')
    
    # Table D
    table_d_data = [
        ['序号', '参数名称', '符号', '单位', '数据要求', '技术说明', '推荐值', '典型范围'],
        ['D1', '运动质量', 'm', 'kg', '折算质量', 'a=(F-F_load)/m', '2.5', '1.5~4.0'],
        ['D2', '弹簧刚度', 'k', 'N/mm', '规格书', '反力曲线斜率', '20', '15~30'],
        ['D3', '线圈匝数', 'N', '匝', '设计值', 'H=NI/L', '1800', '1000~3000'],
        ['D4', '驱动电压', 'U', 'V', '波形图', '电流上升率', 'DC220', 'DC110/220'],
    ]
    create_styled_table(doc, table_d_data, [1.5, 2.0, 1.3, 1.5, 2.0, 2.5, 1.8, 2.0])
    
    doc.add_paragraph()
    
    # ===== SECTION 6 =====
    doc.add_paragraph('6. 数据确认与授权', style='Heading 1')
    
    p = doc.add_paragraph()
    p.add_run('为推进项目进度，若上述数据暂时缺失，甲方可选择以下方式处理：')
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    p = doc.add_paragraph()
    p.add_run('☐ 方式A：由甲方协调供应商补充测试或提供。')
    p.paragraph_format.left_indent = Cm(1.5)
    
    p = doc.add_paragraph()
    p.add_run('☐ 方式B：授权我方依据IEC/GB标准及行业典型数据库选取经验值进行计算。我方不对因原始参数偏差导致的绝对值误差负责。')
    p.paragraph_format.left_indent = Cm(1.5)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    p = doc.add_paragraph()
    run = p.add_run('甲方代表签字：')
    run.bold = True
    p.add_run('____________________')
    
    p = doc.add_paragraph()
    run = p.add_run('日        期：')
    run.bold = True
    p.add_run('____________________')
    
    # Save
    output_path = '/media/large_disk/Projects/Maxwell/Client_Data_Request_Final.docx'
    doc.save(output_path)
    print(f'✓ Document saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    create_document()
